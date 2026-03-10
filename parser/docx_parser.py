from docx import Document


def extract_text_from_docx(file_path):
    """Extract plain text from a .docx file (paragraphs and tables)."""
    text_parts = []
    doc = Document(file_path)

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)
